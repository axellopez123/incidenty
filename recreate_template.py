from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

os.makedirs('app/templates', exist_ok=True)
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Header
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run('Escuela Primaria URBANA 868 "INSTITUTO CABAÑAS"')
run.font.size = Pt(11)
run.bold = True

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('INCIDENCIA ESCOLAR')
run.font.size = Pt(20)
run.bold = True

# Student Info
p = doc.add_paragraph()
p.add_run('Nombre del alumno: ').bold = True
p.add_run('{{student_name}}')
p.add_run(' ' * 10)
p.add_run('Incidencia No. ').bold = True
p.add_run('{{incidencia_id}}')

p = doc.add_paragraph()
p.add_run('Fecha: ').bold = True
p.add_run('{{date}}')
p.add_run(' ' * 20)
p.add_run('Grado/Grupo: ').bold = True
p.add_run('{{grade}} / {{group}}')

# Legal text
p = doc.add_paragraph()
run = p.add_run('Con base en el Reglamento de Conducta para las Escuelas de Educación Básica del Estado de Jalisco, se le informa que el alumno incurrió en una falta de indisciplina:')
run.font.size = Pt(9)

def add_article_section(doc, title_text, numerals, field_name, other_field):
    p = doc.add_paragraph()
    p.add_run(title_text).bold = True
    p = doc.add_paragraph()
    for num in numerals:
        # Checkbox logic using docxtpl syntax
        p.add_run('{% if "' + num + '" in ' + field_name + ' %}[X]{% else %}[ ]{% endif %} ' + num + '   ')
    
    p = doc.add_paragraph()
    p.add_run('{% if ' + other_field + ' %}[X]{% else %}[ ]{% endif %} Otro: {{ ' + other_field + ' or "____________________" }}')

add_article_section(doc, 'ARTÍCULO 13. ACTO LEVE DE INDISCIPLINA', ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII'], 'leve_faction', 'leve_other')
add_article_section(doc, 'ARTÍCULO 14. ACTO GRAVE DE INDISCIPLINA', ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII'], 'grave_faction', 'grave_other')
add_article_section(doc, 'ARTÍCULO 15. ACTO MUY GRAVE DE INDISCIPLINA', ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X', 'XI'], 'muy_grave_faction', 'muy_grave_other')

# Description
doc.add_paragraph('Descripción breve de la situación').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('{{description}}').paragraph_format.space_after = Pt(20)

# Disciplinary
p = doc.add_paragraph()
run = p.add_run('Por lo que estará sujeto (a) a la siguiente medida disciplinaria indicada en el Reglamento de Conducta para las Escuelas de Educación Básica del Estado de Jalisco, artículo 17 según la falta cometida:')
run.font.size = Pt(9)
doc.add_paragraph('{{disciplinary}}').paragraph_format.space_after = Pt(12)

# Agreements
doc.add_paragraph('Acuerdos y compromisos').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('{{acuerdos_compromisos}}').paragraph_format.space_after = Pt(24)

# Signatures
p = doc.add_paragraph('NOMBRE Y FIRMA')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows=2, cols=4)
labels = ['ÁREA QUE RESPONDE DE HOGAR CABAÑAS', 'DOCENTE', 'DIRECTIVO', 'OTRO']
for i, label in enumerate(labels):
    cell = table.cell(1, i)
    cell.text = label
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(7)
    table.cell(0, i).paragraphs[0].add_run('\n\n\n')

doc.save('app/templates/incidencia_template.docx')
